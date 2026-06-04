import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Styles & Colors
# Palette: Deep Navy (#0F2C59), Slate (#4F709C), Dark Neutral (#222831), Light Neutral (#F8F9FA)
COLOR_PRIMARY = RGBColor(15, 44, 89)
COLOR_SECONDARY = RGBColor(79, 112, 156)
COLOR_TEXT = RGBColor(34, 40, 49)
COLOR_MUTED = RGBColor(100, 110, 120)

def set_font(run, name="Calibri", size_pt=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        run = p.add_run(text)
        set_font(run, name="Calibri Light", size_pt=18, bold=True, color=COLOR_PRIMARY)
        # Add a thin bottom border style element if possible, or just space
    elif level == 2:
        run = p.add_run(text)
        set_font(run, name="Calibri", size_pt=14, bold=True, color=COLOR_SECONDARY)
    elif level == 3:
        run = p.add_run(text)
        set_font(run, name="Calibri", size_pt=12, bold=True, italic=True, color=COLOR_TEXT)
    return p

def add_bullet_styled(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, size_pt=11, bold=True, color=COLOR_TEXT)
    r_text = p.add_run(text)
    set_font(r_text, size_pt=11, color=COLOR_TEXT)
    return p

# --- TITLE / COVER ELEMENTS ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(4)
run_title = p_title.add_run("PRODEZK: PERFIL ESTRATÉGICO Y DISEÑO DEL PROGRAMA DE REFERIDOS")
set_font(run_title, name="Calibri Light", size_pt=24, bold=True, color=COLOR_PRIMARY)

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(24)
run_sub = p_sub.add_run("Estructuración integral para la internacionalización y pilar de la estrategia de Partnerships B2B")
set_font(run_sub, name="Calibri", size_pt=12, italic=True, color=COLOR_MUTED)

# Horizontal line separator
p_hr = doc.add_paragraph()
p_hr_border = parse_xml(r'<w:pBdr %s><w:bottom w:val="single" w:sz="12" w:space="1" w:color="0F2C59"/></w:pBdr>' % nsdecls('w'))
p_hr._p.get_or_add_pPr().append(p_hr_border)

# --- SECTION 1 ---
add_heading_styled(doc, "PARTE 1: PERFIL ESTRATÉGICO DE PRODEZK", level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.15
r = p.add_run("1. Resumen Ejecutivo\nProdezk es una firma de consultoría y gestión corporativa especializada en la internacionalización de empresas. Actúa como un puente estratégico para emprendedores locales e internacionales que buscan crear, estructurar y operar negocios legalmente en los Estados Unidos. Con más de 24 años de experiencia en el mercado, la compañía ha evolucionado de ofrecer simples trámites de registro a consolidar un ecosistema integral de más de 100 soluciones corporativas para proteger activos y maximizar el crecimiento empresarial en el mercado norteamericano.")
set_font(r, color=COLOR_TEXT)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.15
r = p.add_run("2. Misión\nSu propósito central es democratizar la creación de empresas en Estados Unidos, simplificando el proceso para que los emprendedores puedan mantener sus negocios al día con las obligaciones legales y federales, potenciando así su crecimiento sin que la complejidad normativa o las barreras del idioma sean un obstáculo.")
set_font(r, color=COLOR_TEXT)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.15
r = p.add_run("3. Visión\nProdezk busca posicionarse como el líder global en la expansión de negocios, con el objetivo de acompañar a miles de empresarios y compañías en la expansión de sus operaciones no solo hacia los Estados Unidos, sino a nivel mundial.")
set_font(r, color=COLOR_TEXT)

add_heading_styled(doc, "4. Valores Corporativos y Filosofía", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("La cultura organizacional y el enfoque de servicio de la empresa se sostienen en cinco pilares (\"De qué están hechos\"):")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, " Creen en el poder de las ideas y la mejora continua de sus procesos.", "• Innovación:")
add_bullet_styled(doc, " Capacidad ágil para tomar decisiones oportunas enfocadas en las necesidades del cliente.", "• Dinamismo:")
add_bullet_styled(doc, " Ejecución transparente y correcta de los procesos, garantizando la máxima seguridad de la información.", "• Integridad:")
add_bullet_styled(doc, " Conexión genuina con los sueños y objetivos comerciales de sus clientes.", "• Pasión:")
add_bullet_styled(doc, " Somos empresarios, como nuestros clientes. Ofrecen un servicio diseñado y prestado por profesionales que entienden sus necesidades, brindando soporte multilingüe (español, inglés y portugués).", "• Cercanía:")

add_heading_styled(doc, "5. Líneas de Negocio y Portafolio de Servicios", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("El modelo de negocio de Prodezk es integral (soluciones end-to-end), dividiendo sus operaciones en las siguientes categorías estratégicas:")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, " Asesoramiento y creación de estructuras jurídicas como LLC, C-CORP y S-CORP en el estado más conveniente para el modelo de negocio.", "• Constitución de Empresas:")
add_bullet_styled(doc, " Asistencia en la apertura de cuentas bancarias comerciales en EE. UU., cuentas en plataformas de pago (Stripe) y cuentas para vendedores de Amazon.", "• Servicios Financieros y E-commerce:")
add_bullet_styled(doc, " Provisión de oficinas virtuales, espacios flexibles, servicio telefónico, agente registrado, enmiendas, apostillas y obtención de certificados (Good Standing).", "• Domicilio y Gestión Corporativa:")
add_bullet_styled(doc, " Gestión integral de impuestos (Income Tax, Sales Tax), tramitación de EIN e ITIN, reportes obligatorios (como el BOI), gestión de nómina y asesorías fiscales.", "• Contabilidad y Cumplimiento Tributario:")

add_heading_styled(doc, "6. Ventajas Competitivas", level=2)
add_bullet_styled(doc, " Pioneros en el sector desde el año 2001, lo que les otorga un profundo know-how del sistema legal y tributario estadounidense.", "• Experiencia Comprobada y Trayectoria:")
add_bullet_styled(doc, " Eliminan la necesidad del cliente de contratar a múltiples proveedores (abogados, contadores, agentes); todo el ciclo de vida de la empresa se gestiona bajo una misma firma.", "• Centralización de Servicios:")
add_bullet_styled(doc, " Sus procesos están diseñados de manera remota y eficiente, permitiendo a extranjeros operar y dirigir sus negocios en EE. UU. desde cualquier parte del mundo, sin necesidad de poseer visa de trabajo o residencia.", "• Apertura sin Fronteras:")
add_bullet_styled(doc, " Tienen su sede principal en el corazón financiero de Miami (Brickell Avenue), complementada con oficinas operativas en Bogotá (Colombia), para facilitar el acceso al mercado latinoamericano.", "• Ubicación Estratégica:")

add_heading_styled(doc, "7. Cifras Clave e Impacto en el Mercado", level=2)
add_bullet_styled(doc, " Más de 24 años de presencia en la industria.", "• Trayectoria:")
add_bullet_styled(doc, " Prestan servicios a empresas de 193 países y abarcan más de 150 sectores económicos diferentes.", "• Alcance:")
add_bullet_styled(doc, " Más de 15,000 empresas creadas con éxito.", "• Comunidad de Clientes:")
add_bullet_styled(doc, " Sus clientes realizan más de 360,000 transacciones comerciales anuales.", "• Volumen Transaccional:")
add_bullet_styled(doc, " Las empresas constituidas a través de Prodezk han logrado facturar más de 85 millones de dólares y generar ganancias superiores a los 16 millones de dólares dentro del ecosistema estadounidense.", "• Impacto Económico:")

# --- SECTION 2 ---
doc.add_page_break()
add_heading_styled(doc, "PARTE 2: DISEÑO ESTRUCTURAL DEL PROGRAMA DE REFERIDOS", level=1)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("El programa de referidos está estructurado para ejecutarse de manera rápida y sencilla, sirviendo como el pilar fundamental del futuro ecosistema de Partnerships B2B de Prodezk. Aprovecha el comportamiento orgánico actual (donde la mayoría de nuevos clientes llegan por recomendación) y lo formaliza a través de un esquema tecnológico escalable y financieramente optimizado basado en el CAC actual del 20% y un ticket promedio de $1,500 - $1,700 USD (base de cálculo: $1,600 USD).")
set_font(r, color=COLOR_TEXT)

add_heading_styled(doc, "1. Estructura de Recompensas: El Modelo Bilateral (Win-Win)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Para eliminar la fricción de la recomendación comercial, se implementa un incentivo de doble vía expresado en valores absolutos líquidos (más atractivos y tangibles psicológicamente que los porcentajes para tickets altos):")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, " Recibe $150 USD en efectivo o crédito a favor para sus servicios recurrentes (Renovación anual, Agente Registrado, Income Tax, reportes BOI). Esto fomenta la retención directa de la base instalada.", "• Para el Referidor (Cliente Actual):")
add_bullet_styled(doc, " Recibe $100 USD de descuento directo aplicable de forma inmediata en su paquete de constitución inicial. Esto destruye la objeción del precio de entrada y asegura el uso del enlace de rastreo.", "• Para el Referido (Nuevo Cliente):")

# Table 1: Estructura de Incentivos Bilaterales
table1 = doc.add_table(rows=3, cols=3)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.autofit = False

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Actor'
hdr_cells[1].text = 'Incentivo Propuesto'
hdr_cells[2].text = 'Justificación Estratégica'

# Format Header row
for cell in hdr_cells:
    shading_elm = parse_xml(r'<w:shd %s w:fill="0F2C59"/>' % nsdecls('w'))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size_pt=10, bold=True, color=RGBColor(255, 255, 255))

row1 = table1.rows[1].cells
row1[0].text = "Referidor (Cliente Actual)"
row1[1].text = "$150 USD en efectivo o crédito corporativo"
row1[2].text = "Fomenta la retención. El crédito se usa en renovaciones, Income Tax o Agente Registrado. Alivia costos futuros."

row2 = table1.rows[2].cells
row2[0].text = "Referido (Nuevo Cliente)"
row2[1].text = "$100 USD de descuento directo"
row2[2].text = "Acelera la conversión inicial y garantiza el uso del enlace de rastreo único para medir la atribución."

for row in table1.rows[1:]:
    for cell in row.cells:
        shading_elm = parse_xml(r'<w:shd %s w:fill="F8F9FA"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                set_font(r, size_pt=9.5, color=COLOR_TEXT)

# Set widths
col_widths = [Inches(1.5), Inches(2.0), Inches(3.0)]
for row in table1.rows:
    for idx, width in enumerate(col_widths):
        row.cells[idx].width = width

add_heading_styled(doc, "2. Reglas de Negocio y Control de Calidad (Términos y Condiciones)", level=2)
add_bullet_styled(doc, "La comisión se genera única y exclusivamente cuando el nuevo cliente realiza el pago total de su trámite y el proceso de incorporación inicia formalmente. No se pagan incentivos por prospectos (leads), únicamente por conversiones reales.", "• Punto de Conversión:")
add_bullet_styled(doc, "Las recompensas en efectivo o créditos se liberan y ven reflejadas en el tablero del referidor 30 días después del pago del referido, mitigando riesgos de fraudes, devoluciones o contratiempos iniciales.", "• Periodo de Maduración/Gracia:")
add_bullet_styled(doc, "No existe un tope máximo de recomendaciones ni de ganancias. Esto permite mapear de forma natural el comportamiento de los usuarios para identificar de forma temprana a los súper-referidores.", "• Volumen Ilimitado:")
add_bullet_styled(doc, "El programa prohíbe explícitamente las auto-referencias (creación de segundas o terceras empresas por el mismo dueño). Estos escenarios comerciales deben ser canalizados mediante estrategias internas de Upselling.", "• Restricción de Auto-referencia:")

add_heading_styled(doc, "3. Infraestructura Tecnológica (Implementación Ágil)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Para una salida a producción ágil y sin fricciones técnicas, se descarta el desarrollo propio de software y se opta por un modelo SaaS (Software as a Service) con integración nativa plug-and-play a la pasarela de pagos actual (Stripe).")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, "Plataformas altamente recomendadas para este modelo. Rewardful destaca por su acople inmediato a Stripe y facilidad en la gestión de servicios corporativos y cobros recurrentes. PartnerStack es la alternativa óptima si se desea integrar desde el inicio una suite preparada para el escalado agresivo hacia canales B2B indirectos.", "• Herramientas Evaluadas (Rewardful / ReferralCandy / PartnerStack):")
add_bullet_styled(doc, "El sistema genera automáticamente enlaces únicos parametrizados (ej: prodezk.com/r/juanperez) para cada cliente actual. Al ingresar el referido, la plataforma instala una cookie de seguimiento. Cuando el sistema procesa el pago a través de Stripe, la herramienta asigna la conversión y el saldo al referidor en piloto automático, notificando a administración para su posterior dispersión o acreditación.", "• Flujo Operativo Automatizado:")

add_heading_styled(doc, "4. Escalabilidad: Transición e Integración al Programa de Partnerships", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("El programa de referidos opera como el nivel primario de captación. Al basarse en un CAC actual del 20% ($320 USD sobre el ticket de $1,600 USD), la fase inicial genera un costo de adquisición consolidado de $250 USD ($100 descuento + $150 comisión), lo que equivale al 15.6% del ticket, arrojando un ahorro inmediato de casi el 5% en utilidad por cliente. Este margen permite estructurar una escala por niveles (Tiers) que incentiva el volumen y capta a aliados comerciales institucionales (contadores, abogados, agencias de marketing en LATAM):")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, "Volumen de 1 a 3 clientes referidos al año. Recibe el incentivo estándar de $150 USD para el referidor y $100 USD para el referido. Costo total de adquisición: $250 USD (15.6% - Ahorro del 4.4% para Prodezk).", "• Nivel 1: Embajador (Cliente Base) —")
add_bullet_styled(doc, "Volumen de 4 a 10 clientes referidos al año. La comisión del referidor incrementa a $170 USD, manteniendo los $100 USD de descuento al nuevo cliente. Costo total de adquisición: $270 USD (16.8% - Ahorro del 3.2% para Prodezk).", "• Nivel 2: Afiliado Activo —")
add_bullet_styled(doc, "Más de 10 clientes referidos al año. Diseñado para firmas aliadas en Latinoamérica. La comisión del referidor sube a $220 USD, y se mantienen los $100 USD para el cliente final. Costo total de adquisición: $320 USD (20% - Equivalente al CAC de pauta digital tradicional).", "• Nivel 3: Partner B2B Oficial —")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Beneficios exclusivos del Nivel Partner: Esquemas avanzados de Revenue Share, acceso a materiales de marca blanca (White-label), co-branding corporativo y la asignación de un ejecutivo de cuenta prioritario dentro de Prodezk. Al igualar el CAC tradicional pero con tráfico orgánico y pre-calificado, la tasa de cierre y el LTV (Lifetime Value) aumentan drásticamente.")
set_font(r, color=COLOR_TEXT)

# Table 2: Niveles del Programa
table2 = doc.add_table(rows=4, cols=5)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.autofit = False

hdr2 = table2.rows[0].cells
hdr2[0].text = 'Nivel'
hdr2[1].text = 'Volumen Anual'
hdr2[2].text = 'Pago Referidor'
hdr2[3].text = 'Bono Nuevo Cliente'
hdr2[4].text = 'Costo Adquisición'

for cell in hdr2:
    shading_elm = parse_xml(r'<w:shd %s w:fill="0F2C59"/>' % nsdecls('w'))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size_pt=10, bold=True, color=RGBColor(255, 255, 255))

data2 = [
    ["Nivel 1: Embajador", "1 a 3 clientes", "$150 USD", "$100 USD", "$250 USD (~15.6%)"],
    ["Nivel 2: Afiliado", "4 a 10 clientes", "$170 USD", "$100 USD", "$270 USD (~16.8%)"],
    ["Nivel 3: Partner B2B", "+10 clientes", "$220 USD", "$100 USD", "$320 USD (20.0%)"]
]

for idx, row_data in enumerate(data2):
    row_cells = table2.rows[idx+1].cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val

for row in table2.rows[1:]:
    for cell in row.cells:
        shading_elm = parse_xml(r'<w:shd %s w:fill="F8F9FA"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                set_font(r, size_pt=9.5, color=COLOR_TEXT)

col_widths2 = [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.4), Inches(1.4)]
for row in table2.rows:
    for idx, width in enumerate(col_widths2):
        row.cells[idx].width = width

add_heading_styled(doc, "5. Kit de Comunicación Estratégica (Activos de Marketing)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Para asegurar una adopción masiva y rápida, se dispondrá un kit de elementos listos para ser utilizados por los clientes en su comunicación diaria:")
set_font(r, color=COLOR_TEXT)

p_box = doc.add_paragraph()
p_box.paragraph_format.left_indent = Inches(0.25)
p_box.paragraph_format.right_indent = Inches(0.25)
p_box.paragraph_format.space_before = Pt(6)
p_box.paragraph_format.space_after = Pt(6)
box_border = parse_xml(r'<w:pBdr %s><w:left w:val="single" w:sz="24" w:space="4" w:color="4F709C"/></w:pBdr>' % nsdecls('w'))
p_box._p.get_or_add_pPr().append(box_border)

r_box = p_box.add_run("Plantilla Controlada para Compartir vía WhatsApp:\n")
set_font(r_box, size_pt=10.5, bold=True, color=COLOR_SECONDARY)
r_box_text = p_box.add_run("\"¡Hola! Te comparto esto porque sé que estás buscando expandir tu negocio. Yo constituí mi empresa en EE. UU. con Prodezk y la verdad el proceso fue súper ágil y en español. Me dieron un enlace de invitación que te regala $100 USD de descuento en la creación de tu empresa. Te lo dejo por aquí por si te sirve: [Insertar Enlace Único]\"")
set_font(r_box_text, size_pt=10.5, italic=True, color=COLOR_TEXT)

add_heading_styled(doc, "6. Roadmap de Implementación Ágil (Tiempos Racionales)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("El plan de despliegue comprende una ventana de tiempo de 6 semanas, garantizando la correcta integración técnica y control del programa:")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, " Establecer los montos exactos y techos de las recompensas, validación del marco de cumplimiento legal de las bonificaciones comerciales y redacción final de los Términos y Condiciones generales. Selección y contratación de la plataforma SaaS (ej. Rewardful).", "• Semana 1: Definición Estratégica y Legal —")
add_bullet_styled(doc, " Integración técnica de la plataforma de referidos elegida con la API de Stripe y sincronización de datos con el CRM corporativo de Prodezk. Automatización y pruebas de flujos de correos transaccionales del sistema.", "• Semana 2: Setup Tecnológico y Conectividad —")
add_bullet_styled(doc, " Diseño y maquetación de la Landing Page informativa (prodezk.com/referidos). Estructuración gráfica del Kit de Referidos dentro del panel de usuario y configuración del pop-up en la página de confirmación de compra.", "• Semana 3: Creación de Activos de Marketing y Contenidos —")
add_bullet_styled(doc, " Simulación integral de flujos de compra en pasarela de pruebas para validar la correcta aplicación del descuento al referido, el correcto despliegue de cookies de rastreo y la correcta asignación del balance en la cuenta del referidor.", "• Semana 4: Pruebas Internas y QA —")
add_bullet_styled(doc, " Lanzamiento Beta cerrado dirigido al Top 10% de los clientes más antiguos y activos de Prodezk. Canalización de feedback inmediato, monitoreo del comportamiento y afinación de posibles desajustes operativos.", "• Semana 5: Lanzamiento Soft (Fase Beta Controlada) —")
add_bullet_styled(doc, " Despliegue masivo mediante campaña de email marketing a toda la base general de clientes activos. Activación de disparadores automáticos (triggers) en el CRM para invitar a todo cliente que cumpla 30 días de servicio exitoso.", "• Semana 6: Lanzamiento Oficial y Automatización General —")

add_heading_styled(doc, "7. Indicadores Clave de Rendimiento (KPIs de Control Financiero)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Para evaluar la efectividad del programa frente a la mesa directiva, se reportarán de forma mensual las siguientes métricas:")
set_font(r, color=COLOR_TEXT)

add_bullet_styled(doc, " Mide el porcentaje total de la base de datos de clientes activos que comparte o interactúa con su enlace único al menos una vez al mes. Meta del primer trimestre: >15%.", "1. Tasa de Participación del Programa:")
add_bullet_styled(doc, " Relación porcentual entre los clics totales registrados en los enlaces de recomendación frente a transacciones pagadas en Stripe. Al ser leads pre-calificados de confianza, la meta esperada es un ratio superior al promedio de pauta digital.", "2. Tasa de Conversión de Referidos:")
add_bullet_styled(doc, " Peso relativo que tienen las ventas originadas por el programa de referidos en comparación con el volumen total de incorporaciones mensuales de Prodezk. Meta primer trimestre: 10% de la facturación global.", "3. Penetración de Ventas Orgánicas por Referencia:")
add_bullet_styled(doc, " Control financiero estricto para garantizar que el costo de adquisición ponderado (descuentos + comisiones + costo de plataforma) se mantenga consistentemente por debajo del techo actual de $320 USD (20% del CAC).", "4. ROI del Programa y Eficiencia del CAC:")

doc.save("Perfil_Estrategico_y_Programa_Referidos_Prodezk.docx")
print("Done")